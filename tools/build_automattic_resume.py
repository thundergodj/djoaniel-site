import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PARSER = argparse.ArgumentParser()
PARSER.add_argument("--target", choices=("automattic", "remote", "rwazi"), default="automattic")
ARGS = PARSER.parse_args()
TARGET = ARGS.target
OUT_NAME = {
    "automattic": "Djoaniel_Hernandez_Automattic_Resume.docx",
    "remote": "Djoaniel_Hernandez_Remote_Resume.docx",
    "rwazi": "Djoaniel_Hernandez_Rwazi_Resume.docx",
}[TARGET]
OUT = ROOT / "output" / "docx" / OUT_NAME

PROFILE_TEXT = {
    "automattic": (
        "Senior Product Designer and UX leader with sixteen years turning complex, high-stakes systems into readable products. "
        "Experience spans private wealth, banking, healthcare, energy, public services, retail, and consumer commerce. "
        "Works end to end—from discovery and interaction models through design systems, engineering partnership, and implementation review. "
        "Uses AI, automation, and code-assisted workflows to accelerate research, explore alternatives, build working prototypes, and keep design rationale close to what ships."
    ),
    "remote": (
        "Senior Product Designer and systems thinker with sixteen years turning complex, high-stakes workflows into clear, scalable products. "
        "Leads ambiguous B2B and financial-platform problems end to end—from research and product framing through interaction design, high-fidelity UI, design systems, engineering partnership, and implementation review. "
        "Experienced in distributed collaboration, persuasive design narratives, accessibility, and reusable platform patterns. "
        "Uses AI and code-assisted prototyping to explore functional flows quickly while keeping product judgement, evidence, and final quality human-owned."
    ),
    "rwazi": (
        "Senior Product Designer and systems thinker with sixteen years turning dense, high-stakes data and workflows into clear enterprise products. "
        "Leads ambiguous B2B and financial-platform problems end to end, from research and product framing through interaction design, data presentation, design systems, engineering partnership, and implementation review. "
        "Experienced in distributed collaboration, business-oriented design narratives, accessibility, and reusable platform patterns. "
        "Uses Claude, Codex, Figma, and working HTML/CSS/JavaScript prototypes to explore AI-assisted experiences while keeping evidence, uncertainty, and final product judgement visible."
    ),
}[TARGET]

CORE_TEXT = {
    "automattic": "Product strategy & discovery · Interaction design · Complex workflows · Design systems · WCAG 2.2 AA · Figma · HTML/CSS/JavaScript · AI-assisted prototyping · User research & synthesis · Workshop facilitation · Async documentation",
    "remote": "Product strategy & discovery · Systems thinking · Complex B2B SaaS workflows · Interaction & UI design · Design systems · WCAG 2.2 AA · Figma · HTML/CSS/JavaScript · AI prototyping · User research & synthesis · Persuasive storytelling · Async collaboration",
    "rwazi": "Product strategy & discovery | Enterprise data visualization | Complex B2B SaaS workflows | AI/LLM interaction design | Design systems | Figma | Claude Code/Codex | HTML/CSS/JavaScript | User research & synthesis | WCAG 2.2 AA | GitHub | Async collaboration",
}[TARGET]

FONT = "Arial"
INK = RGBColor(23, 31, 42)
ACCENT = RGBColor(30, 87, 153)
MUTED = RGBColor(88, 96, 106)
LIGHT = RGBColor(213, 219, 226)


def set_font(run, size=None, bold=None, italic=None, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_geometry(table, widths_dxa, indent_dxa=0):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def create_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "330")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "330")
    ind.set(qn("w:hanging"), "150")
    p_pr.extend([tabs, ind])
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(r_fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(doc, num_id, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3.2)
    p.paragraph_format.line_spacing = 1.04
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    set_font(p.add_run(text), size=9.35)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    set_font(p.add_run(text.upper()), size=10.3, bold=True, color=ACCENT)
    p._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))
    return p


def add_role(doc, title, company, dates, location=None):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.0
    set_font(p.add_run(title), size=10.2, bold=True)
    set_font(p.add_run(f"  |  {company}"), size=10.2, bold=True, color=ACCENT)
    meta = f"{dates}" + (f"  ·  {location}" if location else "")
    set_font(p.add_run(f"\n{meta}"), size=8.55, color=MUTED)
    return p


def add_project(doc, title, descriptor, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3.5)
    p.paragraph_format.space_after = Pt(2.2)
    p.paragraph_format.line_spacing = 1.04
    set_font(p.add_run(title), size=9.5, bold=True)
    set_font(p.add_run(f"  |  {descriptor}\n"), size=8.7, bold=True, color=ACCENT)
    set_font(p.add_run(text), size=9.1)
    return p


def add_hyperlink(paragraph, text, url, color=ACCENT):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    shade = OxmlElement("w:color")
    shade.set(qn("w:val"), str(color))
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    r_pr.extend([r_fonts, size, shade, underline])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.52)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.62)
section.right_margin = Inches(0.62)
section.header_distance = Inches(0.25)
section.footer_distance = Inches(0.25)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(9.35)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.04

for style_name in ("Heading 1", "Heading 2", "Heading 3"):
    style = styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.color.rgb = ACCENT

bullet_num_id = create_bullet_numbering(doc)

# customer_pack-inspired resume header; named overrides: no running chrome,
# no bottom rule, compact contact grid, 0.62-inch margins, Arial throughout.
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(1)
title.paragraph_format.line_spacing = 1.0
set_font(title.add_run("DJOANIEL HERNANDEZ"), size=22, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_before = Pt(0)
subtitle.paragraph_format.space_after = Pt(5)
subtitle.paragraph_format.line_spacing = 1.0
set_font(subtitle.add_run("Senior Product Designer  ·  NN/g Certified UX Expert"), size=11.1, bold=True, color=ACCENT)

contact = doc.add_table(rows=1, cols=2)
set_repeat_table_geometry(contact, [5200, 5235], indent_dxa=0)
left, right = contact.rows[0].cells
lp = left.paragraphs[0]
lp.paragraph_format.space_after = Pt(0)
set_font(lp.add_run("Quezon City, Philippines  ·  +63 917 876 0826  ·  "), size=8.8, color=MUTED)
add_hyperlink(lp, "djoaniel@gmail.com", "mailto:djoaniel@gmail.com", color=MUTED)
rp = right.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rp.paragraph_format.space_after = Pt(0)
add_hyperlink(rp, "djoaniel.com", "https://www.djoaniel.com/")
set_font(rp.add_run("  ·  "), size=8.8, color=MUTED)
add_hyperlink(rp, "LinkedIn", "https://www.linkedin.com/in/djoanielhernandez/")

add_heading(doc, "Profile")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.paragraph_format.line_spacing = 1.06
set_font(
    p.add_run(PROFILE_TEXT),
    size=9.55,
)

skills = doc.add_paragraph()
skills.paragraph_format.space_before = Pt(0)
skills.paragraph_format.space_after = Pt(3)
skills.paragraph_format.line_spacing = 1.05
set_font(skills.add_run("CORE: "), size=8.6, bold=True, color=ACCENT)
set_font(
    skills.add_run(CORE_TEXT),
    size=8.85,
)

add_heading(doc, "Experience")
add_role(doc, "Senior Product Designer", "Avaloq", "June 2021–Present", "Makati / distributed")
for text in (
    "Design private-wealth and banking interfaces where clarity and precision are high-stakes, owning discovery, flows, prototypes, interaction rules, developer handoff, and implementation review across complex B2B modules.",
    "Partner with engineers, business analysts, and product stakeholders across distributed teams to turn ambiguous requirements into deterministic behaviour and surface feasibility issues before build.",
    "Structure reusable form-field and input components in Figma, improving consistency and maintainability across product areas while preserving the context required by specialised financial workflows.",
    "Built the inspection sheet used for WCAG 2.2 AA accessibility work: selected checks across five product surfaces plus an inherited platform pass, making the audit method repeatable instead of deadline-driven.",
    "Facilitate Scrum ceremonies, discovery workshops, story grooming, and prioritisation; mentor designers and present internally on AI tooling, design writing, accessibility, and emerging practice.",
):
    add_bullet(doc, bullet_num_id, text)

if TARGET == "rwazi":
    add_bullet(
        doc,
        bullet_num_id,
        "Led GenAI and innovation initiatives and use Claude, Codex, and working front-end prototypes to explore AI-assisted product behaviour, edge cases, and implementation constraints before build.",
    )

add_role(doc, "UX Associate Manager", "Accenture Liquid Studio", "December 2015–July 2021", "Philippines / Singapore / United States / Japan")
for text in (
    "Led functional and product design for regulated enterprise engagements, translating research, workshops, and business constraints into user journeys, interaction specifications, prototypes, and build-ready stories.",
    "Worked across Merck INTIENT, Florida Power & Light nuclear maintenance tools, Singapore Ministry of Health, and Singapore's National Trade Programme—bridging product owners, business analysts, designers, and engineers.",
    "Designed TESSA's conversational engine and portable skills module for Plan International Philippines, translating ordinary-language answers into employer-ready skills and a downloadable résumé.",
    "Built internal UX capability through team composition, bootcamps, training plans, mentoring, and design-thinking workshops delivered in Tokyo and Singapore.",
):
    add_bullet(doc, bullet_num_id, text)

if TARGET == "rwazi":
    add_bullet(
        doc,
        bullet_num_id,
        "Delivered data-visualization design-thinking training for consultants in Tokyo and designed reporting and decision workflows across regulated healthcare, trade, energy, and banking systems.",
    )

add_role(doc, "Co-founder & Vice President, New Product Development", "SariSari Snaps", "2025–Present", "Consumer product venture")
for text in (
    "Took a rough 3D-printed prototype through material failures, fit tests, colour elimination, and parametric geometry into a nine-format magnetic photo-frame system driven by one shared OpenSCAD model.",
    "Own product rules, prototyping, production constraints, and e-commerce experience; use AI-assisted specification and build workflows while retaining final responsibility for decisions, testing, and quality.",
):
    add_bullet(doc, bullet_num_id, text)

doc.add_page_break()

add_heading(doc, "Selected Product Work")
add_project(
    doc,
    "The site is the work sample",
    "Design system · Front-end · Accessibility",
    "Designed and hand-coded the current portfolio in semantic HTML, CSS, and JavaScript. The type lock, state palette, responsive shell, accessibility behaviours, and working case-study instruments are versioned in GitHub and deployed to Vercel on every push.",
)
add_project(
    doc,
    "Break It Yourself",
    "Enterprise accessibility operations",
    "Converted a European regulatory deadline and WCAG 2.2 AA target into a live, repeatable inspection system. The interactive reconstruction shows how component decisions break or restore specific criteria across tables, forms, product surfaces, and transaction review.",
)
add_project(
    doc,
    "The Art of Nothing",
    "Private-wealth platform · Systems design",
    "Authored one deterministic null-value rule to replace five competing interpretations across a private-wealth platform, coordinating the seven roles needed to move the rule from ambiguity into shipped behaviour.",
)
add_project(
    doc,
    "The Unhappy Path",
    "Retail commerce · Component/state library",
    "Six retail page templates exposed a deeper problem: states such as no size selected, every filter active, or a missing hero image were undefined. Reframed the deliverable around the component-and-state system underneath the screens.",
)
add_project(
    doc,
    "TESSA",
    "Conversation design · Social impact",
    "Designed a Messenger career coach that teaches its own interaction model, asks for consent before data, translates ordinary experience into résumé-ready skills, and hands the user a portable PDF rather than trapping value inside the product.",
)

add_heading(doc, "AI & Code-Assisted Practice")
for text in (
    "Use Claude, ChatGPT/Codex, and code-based prototyping to move between research synthesis, interaction alternatives, specification, implementation, and QA without losing the design rationale between artefacts.",
    "Build working HTML/CSS/JavaScript prototypes and use AI as a critique and production partner: challenge edge cases, inspect accessibility and responsive behaviour, test alternatives, and document why the selected path survived.",
    "Share methods through internal talks, mentoring, and written guidance; automation accelerates exploration and delivery, while product judgement, evidence, accessibility, and final quality remain human-owned.",
):
    add_bullet(doc, bullet_num_id, text)

add_heading(doc, "Leadership & Collaboration")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
p.paragraph_format.line_spacing = 1.05
set_font(
    p.add_run(
        "Design leadership · Cross-functional facilitation · Product-owner partnership · Design–engineering alignment · Distributed delivery · Mentoring · Design critique · Design operations · Agile/Scrum · Jira · Confluence"
    ),
    size=9.1,
)

add_heading(doc, "Earlier Career")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
p.paragraph_format.line_spacing = 1.05
set_font(p.add_run("Web design and creative-services leadership, 2005–2015. "), size=9.25, bold=True)
set_font(
    p.add_run(
        "Designed and implemented web interfaces, e-commerce and email experiences; translated visual design into HTML/CSS; managed multidisciplinary production teams; and built the front-end foundation that still informs current design–engineering collaboration."
    ),
    size=9.25,
)

add_heading(doc, "Certification & Education")
cert = doc.add_paragraph()
cert.paragraph_format.space_after = Pt(2.5)
set_font(cert.add_run("Nielsen Norman Group — Certified UX Expert"), size=9.35, bold=True)
set_font(cert.add_run("  ·  Certificate #1046866  ·  2021"), size=8.9, color=MUTED)

edu = doc.add_paragraph()
edu.paragraph_format.space_after = Pt(2.5)
set_font(edu.add_run("Undergraduate studies"), size=9.35, bold=True)
set_font(
    edu.add_run(
        "  ·  Information Technology, AMA College  ·  Computer Systems Design and Programming, AMA Computer Learning Center"
    ),
    size=8.9,
    color=MUTED,
)

availability = doc.add_paragraph()
availability.paragraph_format.space_before = Pt(6)
availability.paragraph_format.space_after = Pt(0)
availability.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(availability.add_run("Based in Quezon City, Philippines  ·  Available on 30 days' notice"), size=8.6, bold=True, color=ACCENT)

# Remove table borders from the contact grid and suppress header/footer chrome.
tbl_pr = contact._tbl.tblPr
borders = OxmlElement("w:tblBorders")
for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
    border = OxmlElement("w:" + edge)
    border.set(qn("w:val"), "nil")
    borders.append(border)
tbl_pr.append(borders)

doc.core_properties.title = "Djoaniel Hernandez - Senior Product Designer"
subject_target = {"automattic": "Automattic", "remote": "Remote", "rwazi": "Rwazi"}[TARGET]
doc.core_properties.subject = f"Application résumé for {subject_target}"
doc.core_properties.author = "Djoaniel Hernandez"
doc.core_properties.keywords = "Product design, UX, design systems, accessibility, AI, prototyping"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
